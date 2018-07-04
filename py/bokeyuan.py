import urllib.request as ure
import re
import urllib.parse
from delayed import WaitFor
import os
import lxml.html
import docx

def download(url,num=2):
    print('下载中:'+url)
    headers={'User-Agent':'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3325.181 Safari/537.36'}
    request=ure.Request(url,headers=headers)
    try:
        html=ure.urlopen(request).read()
    except ure.URLError as e:
        print('下载失败：'+e.reason)
        html=None
        if num>0:
            if hasattr(e,'code') and 500<=e.code<600:
                return download(url,num-1)
    return html
def link_crawler(seed_url,link_regex):
    html =download(seed_url)
    crawl_queue=[]
    for link in get_links(html):
        if re.match(link_regex,link):
            link=urllib.parse.urljoin(seed_url,link)
            if link not in crawl_queue:
                crawl_queue.append(link)
    x=0
    waitFor = WaitFor(2)
    while crawl_queue:
        url = crawl_queue.pop()
        waitFor.wait(url)
        html = download(url)
        tree = lxml.html.fromstring(html)
        #title = tree.xpath('//a[@id="cb_post_title_url"]')
        title = tree.xpath('//*[@id="j_core_title_wrap"]/h3')
        #content = tree.xpath('//div[@id="cnblogs_post_body"]/p')
        content = tree.xpath('//*[@class="d_post_content j_d_post_content "]')

        #pre = tree.xpath('//pre')
        #img = tree.xpath('//div[@id="cnblogs_post_body"]/p/img/@src')
        img = tree.xpath('//*[@class="d_post_content j_d_post_content "]/img')
        os.chdir(r'C:\Users\lenovo\Desktop\py\美图吧')
        doc = docx.Document()
        try:
            doc.add_heading(title[0].text_content(), 0)
        except IndexError as e:
            continue
        for i in content:
            doc.add_paragraph(i.text_content())
        #for p in pre:
            #doc.add_paragraph(p.text_content())
        for i in img:
            try:
                ure.urlretrieve(i, '0.jpg')
                doc.add_picture('0.jpg')
            except:
                print("无法获取图片")
        filename = title[0].text_content()[:8] + '.docx'
        if str(filename) in os.listdir(r'C:\Users\lenovo\Desktop\py\美图吧'):
           doc.save(title[0].text_content()[:8] + str(x) + '.docx')
           x += 1
        else:
            doc.save(filename)
def get_links(html):
   webpage_regex = re.compile('<a[^>]+href=["\'](.*?)["\']',re.IGNORECASE)
   html = html.decode('ISO-8859-1')
   return webpage_regex.findall(html)
link_crawler('https://tieba.baidu.com/f?ie=utf-8&kw=%E7%BE%8E%E5%9B%BE&fr=search','.*/tieba.baidu.com/p/.*')
